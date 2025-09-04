from openai import OpenAI       # Cliente oficial de OpenAI (nuevo cliente python).
import streamlit as st         # Streamlit para la interfaz web.
from docx import Document     # python-docx para crear documentos .docx
from io import BytesIO        # Para guardar el .docx en memoria (no en disco)
import os                     # para acceder a variables de entorno (API KEY)
import pandas as pd


# -----------------------
# Función: generar_articulo
# -----------------------
def generar_articulo(client, topic):
    """
    Llama a la API de OpenAI (Chat Completions) para generar un artículo SEO sobre `topic`.
    - client: instancia de openai.OpenAI
    - topic: texto (string) con el tema del artículo

    Devuelve: texto (string) con el artículo, o lanza excepción si falla.
    """
    try:
       
        completion = client.chat.completions.create(
            model="gpt-4o",   # modelo; ajusta según tu acceso (p. ej. "gpt-4o-mini" si lo prefieres)
            messages=[
                {"role": "system", "content": "Eres un experto en redacción de artículos SEO."},
                {"role": "user", "content": f"Escribe un artículo optimizado para SEO sobre: {topic}"}
            ],
            max_tokens=1000   # controla la longitud de la respuesta
        )

        article = completion.choices[0].message.content.strip()
        return article

    except Exception as e:
        # No mezclar UI y lógica si quieres reusar la función fuera de Streamlit.
        # Aquí devolvemos la excepción hacia arriba para que la UI la gestione.
        raise

# -----------------------
# Función: create_word_doc
# -----------------------
def generar_doc(article_text):
    """
    Crea un archivo .docx en memoria con el texto `article_text`.
    Devuelve un objeto BytesIO listo para ser enviado como descarga.
    """
    doc = Document()                          # crea documento vacío
    doc.add_heading('Artículo Generado', level=0)  # título principal
    # Si quieres encabezados más ricos (título, subtítulos), podrías parsear el artículo y añadir varios párrafos/estilos.
    doc.add_paragraph(article_text)           # agrega el contenido del artículo como un párrafo

    buffer = BytesIO()                        # buffer en memoria
    doc.save(buffer)                          # guarda el .docx en el buffer
    buffer.seek(0)                            # vuelve al inicio del buffer para poder leerlo
    return buffer

# -----------------------
# Función: generar_codigo (placeholder / ejemplo)
# -----------------------
def generar_codigo(client, prompt):
    """
    Ejemplo de función que pediría al modelo que genere código.
    Aquí devolvemos el texto directamente; puedes hacer post-procesamiento si quieres.
    """
    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Eres un programdor de Python genera codigo correctamente formateado, sin explicacion ni comentarios adicionales."},
                {"role": "user", "content": f"Crea un codigo de python para {prompt}. Imiportante: Proporciona: SOLO el codigo, sin ningun texto explicativo"}
            ],
            temperature=0.2,
            max_tokens=800
        )
        code =completion.choices[0].message.content.strip()
        return code
    except Exception as e:
        raise

# -----------------------
# Función: generar_tabla (placeholder / ejemplo)
# -----------------------
def generar_tabla(client, topic):
    """
    Ejemplo: pedirle al modelo que devuelva una tabla en formato CSV o Markdown
    que luego puedas convertir a pandas.DataFrame.
    """
    try:
        prompt = f"Genera una tabla en formato CSV  sobre: {topic} los dats deben ser coherentes y bien estructurados"
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Eres un experto en estructurar datos tabulares."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=400
        )
        csv_text = completion.choices[0].message.content.strip()
        return csv_text
    except Exception as e:
        raise
    # -----------------------
# Función: create_word_doc
# -----------------------



def generar_excel(csv_text):
    '''Convertir un texto CSV en un archivo Excel guardado en un buffer en memoria'''
    
    # 1. Dividir el texto CSV en líneas, eliminando espacios en blanco y líneas vacías
    lines = [line.strip() for line in csv_text.split('\n') if line.strip()]
    
    # 2. Si no hay líneas con datos, lanzar un error
    if not lines:
        raise ValueError("No hay datos para procesar")
    
    # 3. La primera línea contiene los encabezados (nombres de columnas), los separo por coma
    headers = lines[0].split(',')
    
    # 4. El resto de las líneas son filas de datos, las separo también por coma para cada columna
    data_rows = [line.split(',') for line in lines[1:]]
    
    # 5. Creo un DataFrame de pandas con los datos y los encabezados como columnas
    df = pd.DataFrame(data_rows, columns=headers)
    
    # 6. Creo un buffer en memoria para guardar el archivo Excel sin escribir en disco
    buffer = BytesIO()
    
    # 7. Abro un escritor de Excel que escribirá en el buffer, usando el motor 'xlsxwriter'
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        
        # 8. Escribo el DataFrame en una hoja llamada 'Datos', sin incluir índices
        df.to_excel(writer, index=False, sheet_name='Datos')
        
        # 9. Obtengo la hoja de Excel para poder modificarla (por ejemplo el ancho de columnas)
        worksheet = writer.sheets['Datos']
        
        # 10. Ajusto automáticamente el ancho de cada columna según el contenido
        for idx, col in enumerate(df.columns):
            # Calculo la longitud máxima entre los datos y el nombre de columna
            max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
            # Establezco el ancho de la columna en la hoja Excel
            worksheet.set_column(idx, idx, max_len)
    
    # 11. Vuelvo al inicio del buffer para que pueda leerse desde el principio
    buffer.seek(0)
    
    # 12. Retorno el buffer que contiene el archivo Excel generado en memoria
    return buffer,df
